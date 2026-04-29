"""
文档处理器 — PDF / TXT / MD / DOCX 加载与分块
"""
import io

import structlog
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from pypdf import PdfReader
from docx import Document as DocxDocument

from config import settings

logger = structlog.get_logger()


class DocProcessor:
    def __init__(self, rag_config=None):
        self.rag_config = rag_config or settings.rag
        self.chunk_size = self.rag_config.chunk_size
        self.chunk_overlap = self.rag_config.chunk_overlap
        self.splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            separators=["\n\n", "\n", "。", "！", "？", " ", ""],
            length_function=len,
        )
    async def process(self, file_bytes: bytes, file_type: str, file_name: str) -> list[Document]:
        """
        加载文档 → 分块 → 添加元数据
        """
        try:
            docs = self._load_documents(file_bytes, file_type, file_name)

            chunks = self.splitter.split_documents(docs)

            for i, chunk in enumerate(chunks):
                chunk.metadata.update({
                    "source": file_name,
                    "chunk_index": i,
                    "file_type": file_type,
                    "total_chunks": len(chunks),
                })

            logger.info("doc.processed", file_name=file_name, chunks=len(chunks))
            return chunks

        except Exception as e:
            logger.error("doc.process.failed", file_name=file_name, error=str(e))
            return []

    def _load_documents(self, file_bytes: bytes, file_type: str, file_name: str) -> list[Document]:
        if file_type in {"txt", "md"}:
            text = file_bytes.decode("utf-8", errors="replace")
            return [Document(page_content=text, metadata={"source": file_name, "file_type": file_type})]

        if file_type == "pdf":
            reader = PdfReader(io.BytesIO(file_bytes))
            docs: list[Document] = []
            for page_index, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                if text.strip():
                    docs.append(
                        Document(
                            page_content=text,
                            metadata={"source": file_name, "file_type": file_type, "page": page_index + 1},
                        )
                    )
            return docs

        if file_type == "docx":
            doc = DocxDocument(io.BytesIO(file_bytes))
            paragraphs = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
            text = "\n".join(paragraphs)
            return [Document(page_content=text, metadata={"source": file_name, "file_type": file_type})] if text else []

        raise ValueError(f"Unsupported file type: {file_type}")
